"""Document Service: Handles parsing of CVs and generation of research reports."""

import pathlib
import tempfile
from io import BytesIO
from datetime import datetime, timezone
from typing import List, Union, Optional, Any

from markitdown import MarkItDown
from docx import Document
from src.state import AgentState


class DocumentService:
    """
    Orchestrates file conversions and document generation.
    Abstracts MarkItDown and python-docx complexities.
    """

    def __init__(self, output_path: str = "files/output"):
        self.md = MarkItDown()
        self.output_dir = pathlib.Path(output_path)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def convert_to_text(self, file_content: bytes, original_name: str) -> str:
        """Converts uploaded bytes into clean text using a temporary file."""
        suffix = pathlib.Path(original_name).suffix
        with tempfile.NamedTemporaryFile(delete=True, suffix=suffix) as tmp:
            tmp.write(file_content)
            tmp.flush()
            result = self.md.convert(tmp.name)
            return result.text_content

    def ingest_directory(self, folder_path: str) -> str:
        """Aggregates all documents in a folder into a single context string."""
        path = pathlib.Path(folder_path)
        if not path.exists():
            return ""

        fragments = []
        for file_path in path.iterdir():
            if file_path.is_file():
                try:
                    result = self.md.convert(str(file_path))
                    fragments.append(
                        f"--- SOURCE: {file_path.name} ---\n{result.text_content}"
                    )
                except Exception as e:
                    print(f"Failed to ingest {file_path.name}: {e}")

        return "\n\n".join(fragments)

    def generate_research_report(self, state: AgentState) -> BytesIO:
        """
        Constructs a professional .docx report from the agent state.
        Returns a BytesIO buffer instead of writing to disk directly.
        """
        writer_data = state.get("writer_data")
        candidate_name = (
            state.get("cv_data").full_name if state.get("cv_data") else "Candidate"
        )

        doc = Document()
        doc.add_heading(f"Job Research Report: {candidate_name}", 0)

        if not writer_data or not writer_data.jobs:
            doc.add_paragraph("No job matches found in this session.")
        else:
            for job in writer_data.jobs:
                self._build_job_section(doc, job)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _build_job_section(self, doc: Document, job: Any):
        """Private helper to maintain clean Word formatting."""
        doc.add_heading(job.title, level=1)
        doc.add_paragraph(f"Company: {job.company}", style="Caption")
        doc.add_paragraph(f"Match Score: {job.top_applicant_score}%")
        doc.add_paragraph(f"URL: {job.job_url}")

        if job.salary_min or job.salary_max:
            salary = f"£{job.salary_min or '?'}-{job.salary_max or '?'}"
            doc.add_paragraph(f"Estimated Salary: {salary}")

        doc.add_heading("Analysis & Reasoning", level=2)
        doc.add_paragraph(job.top_applicant_reasoning)

        if job.key_skills:
            doc.add_heading("Targeted Skills", level=2)
            for skill in job.key_skills:
                doc.add_paragraph(skill, style="List Bullet")

        doc.add_page_break()

    def get_standard_filename(self, candidate_name: str) -> str:
        """Generates a standardized filename with timestamp."""
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M")
        name = candidate_name.replace(" ", "_").lower()
        return f"{ts}_{name}_job_research.docx"
