from glob import glob
from os import path
from datetime import datetime
import pathlib

from markitdown import MarkItDown
from docx import Document

from src.state import AgentState
from src.utils.func import log_message

def ingest_input_folder(folder_path="files/input"):
    """Reads all supported files and returns a single concatenated string."""
    md = MarkItDown()
    aggregated_text = []

    files = glob(path.join(folder_path, "*.*"))

    for file_path in files:
        # MarkItDown automatically detects .docx, .pdf, .pptx, .xlsx, etc.
        try:
            log_message(f"Ingesting with MarkItDown: {path.basename(file_path)}")
            result = md.convert(file_path)
            content = result.text_content
            aggregated_text.append(
                f"--- CONTENT FROM {path.basename(file_path)} ---\n{content}"
            )
        except Exception as e:
            log_message(f"Failed to convert {file_path}: {e}")

    return "\n\n".join(aggregated_text)

def save_findings_to_docx(state: AgentState) -> str:
    """
    Saves a list of analysed jobs and relevant details to a Microsoft Word (.docx) file.

    Args:
        analysed_jobs_json (str): A JSON string containing the job matches.
            The JSON should follow the structure of AnalysedJobMatchList.
    """

    try:
        candidate = state["cv_data"].full_name.replace(" ", "_").lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")

        analysed_jobs = state["writer_data"]
        file_name = f"{timestamp}_{candidate}_job_research.docx"
        log_message(f"Writing Analysis to {file_name}")
        doc = Document()
        doc.add_heading("Top Job Matches & Analysis", 0)

        for job in analysed_jobs.jobs:
            doc.add_heading(job.title, level=1)
            doc.add_paragraph(f"Company: {job.company}")
            doc.add_paragraph(f"Link: {job.job_url}")

            salary_string = []
            if hasattr(job, "salary_min") and job.salary_min:
                salary_string.append(str(job.salary_min))
            if hasattr(job, "salary_max") and job.salary_max:
                salary_string.append(str(job.salary_max))

            if salary_string:
                formatted_salary = " - £".join(salary_string)
                doc.add_paragraph(f"Salary: £{formatted_salary}")

            doc.add_paragraph(f"Summary of role: {job.job_summary}")
            doc.add_paragraph(f"Match Score: {job.top_applicant_score}%")

            doc.add_heading("Why you match:", level=2)
            doc.add_paragraph(job.top_applicant_reasoning)

            doc.add_heading("Tech Stack to Highlight:", level=2)
            if job.tech_stack:
                for tech in job.tech_stack:
                    doc.add_paragraph(tech, style="List Bullet")
            else:
                doc.add_paragraph("No specific tech stack highlighted.")

            doc.add_page_break()

        files_dir = pathlib.Path("files/output")
        files_dir.mkdir(parents=True, exist_ok=True)
        file_path = files_dir / file_name
        doc.save(file_path)

        return f"Successfully saved findings to {file_path}"

    except Exception as e:
        return f"Error saving docx: {str(e)}"
