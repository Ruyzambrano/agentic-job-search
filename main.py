from os import environ as ENV, path
import json
from datetime import datetime
import pathlib
from glob import glob

from docx import Document
from markitdown import MarkItDown
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.text import Text
from rich import box
from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.messages import HumanMessage

from src.state import AgentState
from src.agents.cv_parser import create_cv_parser_agent
from src.agents.researcher import create_researcher_agent
from src.agents.writer import create_writer_agent
from src.graph import create_workflow


def ingest_input_folder(folder_path="files/input"):
    """Reads all supported files and returns a single concatenated string."""
    md = MarkItDown()
    aggregated_text = []

    files = glob(path.join(folder_path, "*.*"))

    for file_path in files:
        # MarkItDown automatically detects .docx, .pdf, .pptx, .xlsx, etc.
        try:
            print(f"Ingesting with MarkItDown: {path.basename(file_path)}")
            result = md.convert(file_path)
            content = result.text_content
            aggregated_text.append(
                f"--- CONTENT FROM {path.basename(file_path)} ---\n{content}"
            )
        except Exception as e:
            print(f"Failed to convert {file_path}: {e}")

    return "\n\n".join(aggregated_text)


def save_findings_to_docx(state: AgentState):
    """
    Saves a list of analysed jobs and relevant details to a Microsoft Word (.docx) file.

    Args:
        analysed_jobs_json (str): A JSON string containing the job matches.
            The JSON should follow the structure of AnalysedJobMatchList.
    """

    try:
        candidate = state["cv_data"].full_name.replace(" ", "_").lower()
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")

        analysed_jobs = state["writer_data"]
        file_name = f"{timestamp}_{candidate}_job_research.docx"
        print(f"LOG: Writing Analysis to {file_name}")
        doc = Document()
        doc.add_heading("Top Job Matches & Analysis", 0)

        for job in analysed_jobs.jobs:
            doc.add_heading(job.title, level=1)
            doc.add_paragraph(f"Company: {job.company}")
            doc.add_paragraph(f"Link: {job.job_url}")

            salary_string = []
            if hasattr(job, "salary_min") and job.salary_min:
                salary_string.append(str(job.salary_min))
            if hasattr(job, "salary_max") and job.salary_max:
                salary_string.append(str(job.salary_max))

            if salary_string:
                formatted_salary = " - £".join(salary_string)
                doc.add_paragraph(f"Salary: £{formatted_salary}")

            doc.add_paragraph(f"Summary of role: {job.job_summary}")
            doc.add_paragraph(f"Match Score: {job.top_applicant_score}%")

            doc.add_heading("Why you match:", level=2)
            doc.add_paragraph(job.top_applicant_reasoning)

            doc.add_heading("Tech Stack to Highlight:", level=2)
            if job.tech_stack:
                for tech in job.tech_stack:
                    doc.add_paragraph(tech, style="List Bullet")
            else:
                doc.add_paragraph("No specific tech stack highlighted.")

            doc.add_page_break()

        files_dir = pathlib.Path("files/output")
        files_dir.mkdir(parents=True, exist_ok=True)
        file_path = files_dir / file_name
        doc.save(file_path)

        return f"Successfully saved findings to {file_path}"

    except Exception as e:
        return f"Error saving docx: {str(e)}"


def pretty_print_jobs_with_rich(json_string):
    console = Console()
    try:
        data = json.loads(json_string)
    except json.JSONDecodeError:
        console.print("[bold red]Error:[/bold red] Invalid JSON string.")
        return

    if not data:
        console.print("[bold red]Error:[/bold red] JSON string shape is wrong:")
        console.print_json(data)

    console.print(
        Panel(
            Text(
                "JOB LISTINGS & APPLICANT INSIGHTS", justify="center", style="bold cyan"
            ),
            box=box.DOUBLE_EDGE,
        )
    )

    for job in data.get("jobs", []):
        table = Table(
            title=f"[bold magenta]{job['title']}[/bold magenta]",
            title_justify="left",
            box=box.ROUNDED,
            expand=True,
        )

        table.add_column("Attribute", style="cyan", width=15)
        table.add_column("Details", style="white")

        table.add_row("Company", job["company"])
        table.add_row("Location", f"{job['location']} ({job['office_days']})")
        salary_parts = [
            f"{s:,}"
            for s in [job.get("salary_min"), job.get("salary_max")]
            if s is not None
        ]
        salary_display = (
            f"£{' - £'.join(salary_parts)}" if salary_parts else "Not Specified"
        )
        table.add_row("Salary", salary_display)
        tech_display = (
            ", ".join(job["tech_stack"])
            if isinstance(job["tech_stack"], list)
            else "N/A"
        )
        table.add_row("Tech Stack", tech_display)
        attributes_display = (
            " • ".join(job["attributes"])
            if isinstance(job["attributes"], list)
            else "N/A"
        )
        table.add_row("Attributes", attributes_display)
        link_text = f"[link={job['job_url']}]Click to view job listing[/link]"
        table.add_row("Listing", link_text)

        console.print(table)

        # Applicant Scoring Section
        score = job["top_applicant_score"]
        score_color = "green" if score >= 85 else "yellow" if score >= 70 else "red"

        score_text = Text()
        score_text.append("\nTOP APPLICANT MATCH: ", style="bold")
        score_text.append(f"{score}%", style=f"bold {score_color}")

        console.print(score_text)
        console.print(
            Panel(
                job["top_applicant_reasoning"],
                title="Match Reasoning",
                title_align="left",
                border_style=score_color,
            )
        )

        console.print("\n")


def get_llm_model(model: str) -> ChatGoogleGenerativeAI:
    return ChatGoogleGenerativeAI(model=model)


def run_job_matcher(raw_context, config: dict):
    cv_parser_agent = create_cv_parser_agent(
        get_llm_model(
            model=ENV.get("CV_PARSE_GEMINI_MODEL"),
        )
    )
    researcher_agent = create_researcher_agent(
        get_llm_model(ENV.get("SEARCH_GEMINI_MODEL"))
    )
    writer_agent = create_writer_agent(get_llm_model(ENV.get("WRITER_GEMINI_MODEL")))

    app = create_workflow(cv_parser_agent, researcher_agent, writer_agent)
    
    desired_job = config.get("configurable", {}).get("role")
    if desired_job:
        desired_job = f"focused on {desired_job}"
    desired_location = config.get("configurable", {}).get("location")
    if desired_location:
        desired_location = f"in {desired_location}"
    content = f"Parse the provided raw cv text and find the best job matches {desired_job} {desired_location}: {raw_context}"
    state = {"messages": [HumanMessage(content=content)]}

    print("----------- Starting Workflow -----------")

    state = app.invoke(input=state, config=config)

    print("\n" + "=" * 30)

    print(save_findings_to_docx(state))
    pretty_print_jobs_with_rich(state["writer_data"].model_dump_json())
    print("SUCCESS: WORKFLOW COMPLETE")


if __name__ == "__main__":
    load_dotenv()
    desired_job = input("What job role are you looking for?\n").strip()
    desired_location = input("Where are you looking today?\n").strip()
    config = {"configurable": {"user_id": "Ruy001", "location": desired_location, "role": desired_job}}
    raw_context = ingest_input_folder("files/input")
    run_job_matcher(raw_context, config)
